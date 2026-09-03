#!/usr/bin/env python3
"""extract_template.py — 模板截取（起止定位 + 原样切片）

不做骨架提取：在招标文件 docx 中确定「响应/投标文件格式」模板的起始与终止
标题，把两者之间的内容**原样截取**为 template.docx（复制源 Document 后删除
范围之外的正文块 → 100% 继承源文档样式、表格、占位符、页眉页脚）。
另产出 template.json 元数据（起止位置、章内标题清单、近似页码、require 命中）。

CLI:
  python extract_template.py <招标文件.docx> -o template.docx [-m template.json]
      [--start 响应文件要求及格式] [--end 价格调整] \
      [--require 投标函,报价表,偏离表,资格审查]

起始定位：候选=标题段中命中 --start（或默认 START_RE）者；多候选时按
「require 命中数 → 内容跨度」打分选最优（避免命中目录引言句）。跨度 < 400 字
或无候选 → 退出码 3，必须向用户确认起止章节。
退出码：0 成功；2 文件错误；3 起止定位失败或 require miss（停下问用户）。
"""
import argparse, json, re, sys
from datetime import datetime

try:
    from docx import Document
    from docx.text.paragraph import Paragraph as DocxPara
except ImportError:
    print("ERROR: python-docx 未安装", file=sys.stderr)
    sys.exit(2)

NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
CN = "零一二三四五六七八九十百千"
CHAPTER_RE = re.compile(r"^\s*第\s*[" + CN + r"0-9１-９]+\s*[章节部分篇卷]")
CN_PREFIX_RE = re.compile(r"^\s*[" + CN + r"]{1,3}\s*[、.．]\s*\S")
NUM_PREFIX_RE = re.compile(r"^(?:\d{1,2}(?:[.．]\d{1,2}){0,3})[、.．\s]\s*\S")
HEADING_STYLE_RE = re.compile(r"(?:Heading|标题)\s*([1-9])", re.I)
START_RE = re.compile(r"(响应|投标|应答)\s*文件.{0,8}(要求及|及|的)?格式|文件格式及附件")
MIN_SPAN_CHARS = 400


def classify(t):
    """(level, family) 或 None。family: chapter/cn/num。"""
    if not t or len(t) > 40 or t[-1] in "。；，;,.":
        return None
    if CHAPTER_RE.match(t):
        return (1, "chapter")
    if CN_PREFIX_RE.match(t):
        return (2, "cn")
    m = re.match(r"^(\d{1,2}(?:[.．]\d{1,2})*)[、.．\s]", t)
    if m and len(t) <= 40:
        dots = m.group(1).count(".") + m.group(1).count("．")
        return (min(dots + 2, 5), "num")
    return None


def is_heading_text(t):
    return classify(t) is not None


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("tender")
    ap.add_argument("-o", "--out", default="template.docx")
    ap.add_argument("-m", "--meta", default="template.json")
    ap.add_argument("--start", default="", help="起始标题关键词（含于标题即命中）")
    ap.add_argument("--end", default="", help="终止标题关键词（命中段本身不含）")
    ap.add_argument("--require", default="")
    a = ap.parse_args()
    req = [r.strip() for r in a.require.split(",") if r.strip()]

    try:
        doc = Document(a.tender)
    except Exception as e:
        print(f"ERROR: 无法打开 {a.tender}: {e}", file=sys.stderr)
        sys.exit(2)

    body = doc.element.body
    children = list(body)
    paras = []  # (child_idx, text, kind, style_heading)
    for i, ch in enumerate(children):
        if ch.tag == NS + "p":
            p = DocxPara(ch, doc)
            t = p.text.strip()
            style_h = bool(HEADING_STYLE_RE.search(p.style.name or ""))
            paras.append((i, t, classify(t) if t else None, style_h))
    if not paras:
        print("ERROR 文档无段落", file=sys.stderr)
        sys.exit(2)
    texts = [(ci, t) for ci, t, k, sh in paras if t]
    total = len(children)

    def span_end(ci, kind):
        """起始之后的终止 child_index：第一个同族同级/更高级标题；chapter 起始见下一 chapter；否则文末。"""
        if a.end:
            for ci2, t2 in texts:
                if ci2 > ci and a.end in t2:
                    return ci2, t2
            return None, None
        lv, fam = kind if kind else (1, "chapter")
        for ci2, t2, k2, sh2 in paras:
            if ci2 <= ci or not t2:
                continue
            k = k2
            if k is None:
                if sh2 and CHAPTER_RE.match(t2):
                    k = (1, "chapter")
                else:
                    continue
            if fam == "chapter" and k[1] == "chapter":
                return ci2, t2
            if k[1] == fam and k[0] <= lv:
                return ci2, t2
        return total, "(至文末)"

    # ---- 起始候选打分 ----
    cands = []
    for ci, t, k, sh in paras:
        if not t or not (k or sh):
            continue
        hit = (a.start in t) if a.start else bool(START_RE.search(t))
        if not hit:
            continue
        e_ci, e_t = span_end(ci, k)
        if e_ci is None:
            print(f"MISS 终止关键词「{a.end}」在候选起始《{t}》之后未命中 —— 向用户确认", file=sys.stderr)
            sys.exit(3)
        body_txt = "".join(x for c2, x in texts if ci <= c2 < e_ci)
        span = len(body_txt)
        hits = sum(1 for r in req if r in body_txt)
        cands.append((hits, span, ci, t, e_ci, e_t))
    if not cands:
        print(f"MISS 未定位到模板起始{'（'+a.start+'）' if a.start else '章（响应/投标文件格式）'}"
              f" —— 必须向用户确认起始/终止章节后再执行", file=sys.stderr)
        sys.exit(3)
    cands.sort(key=lambda x: (-x[0], -x[1], x[2]))
    hits, span, start_ci, start_title, end_ci, end_title = cands[0]
    if span < MIN_SPAN_CHARS:
        print(f"MISS 最优起始候选《{start_title}》跨度仅 {span} 字（<{MIN_SPAN_CHARS}），"
              f"疑似命中目录/引言句 —— 必须向用户确认起止章节（候选：" +
              "；".join(f"《{c[3]}》{c[1]}字" for c in cands[:4]) + "）", file=sys.stderr)
        sys.exit(3)

    # ---- 原样切片：删除 [start_ci, end_ci) 之外的所有顶层正文块 ----
    # 注意：不能只删 w:p / w:tbl —— 招标文件的表格/段落常被 <w:sdt>（内容控件）
    # 或 <w:customXml> 包裹，这些块若不处理会残留在切片开头（"前几页重叠"现象）。
    # 策略：除 body 末尾的 sectPr（节属性，必须保留）外，范围外的块一律整体删除；
    # 范围外的 sdt 整体删除是安全的（其内部内容不可能跨越起始标题）。
    removed = 0
    last_sectpr = body.find(NS + "sectPr")  # body 直接子级里的最终节属性
    for i, ch in enumerate(children):
        if ch is last_sectpr:
            continue
        if start_ci <= i < end_ci:
            continue
        body.remove(ch)
        removed += 1
    doc.save(a.out)

    # ---- 元数据 ----
    tpl = Document(a.out)
    headings, texts2 = [], []
    for p in tpl.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        texts2.append(t)
        if is_heading_text(t) or HEADING_STYLE_RE.search(p.style.name or ""):
            headings.append(t)
    joined = "\n".join(texts2)
    misses = [{"requirement": r, "hit": False} for r in req if r not in joined]
    chars_before = sum(len(x) for c2, x in texts if c2 < start_ci)
    p_start = chars_before // 780 + 1
    p_span = max(1, round(span / 780))
    meta = {
        "source": a.tender, "template": a.out,
        "extracted_at": datetime.now().isoformat(timespec="seconds"),
        "start": {"child_index": start_ci, "title": start_title, "approx_page": p_start},
        "end": {"child_index": end_ci, "title": end_title,
                "approx_page": p_start + p_span - 1},
        "approx_pages": p_span,
        "headings": headings,
        "paragraphs": len(texts2), "tables": len(tpl.tables),
        "removed_blocks": removed, "misses": misses,
        "candidates": [{"title": c[3], "span_chars": c[1], "require_hits": c[0]} for c in cands[:6]],
    }
    with open(a.meta, "w", encoding="utf-8") as f:
        json.dump(meta, f, ensure_ascii=False, indent=1)
    print(f"OK start=「{start_title[:24]}」≈p{p_start} end=「{end_title[:24]}」"
          f" 跨度≈{p_span}页 字符={span} require命中={hits}/{len(req)} "
          f"headings={len(headings)} tables={meta['tables']} misses={len(misses)}")
    print(f"   template: {a.out}  meta: {a.meta}")
    if misses:
        for m in misses:
            print(f"MISS 切片内未找到《{m['requirement']}》—— 向用户确认替代方案后再继续", file=sys.stderr)
        sys.exit(3)


if __name__ == "__main__":
    main()
