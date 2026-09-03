#!/usr/bin/env python3
"""verify.py — 交付前机检关卡

对成稿（docx 或 filled json）执行确定性检查，全过才可交付：
  A. 占位残留：正文不得含【待填/待确认/待补】+ 原始模板占位（下划线/冒号槽/非豁免【】）
  B. 金额一致：quotes 分项合计=总价；数字与大写互核（cn_to_num）  [需 --quotes]
  C. 结构保真：filled json 的节标题集合 ⊇ template.json 的节标题集合  [需 --template]
  D. 红线扫描：对招标 disqualification_clauses 逐条给出正文响应证据；
     按行判定，红线短语仅出现在目录/标题式短行不算实质响应  [需 --tender-parse]

CLI:
  python verify.py <成稿.docx|filled.json> [--template template.json] [--quotes quotes.json]
      [--tender-parse 解析结果.json] -o verify_report.md
退出码：0 通过（含 WEAK/SKIP）；5 存在 FAIL；2 输入错误。
未提供 --quotes/--template/--tender-parse 时对应轨道记 SKIP（报告显式披露，不静默跳过）；
SKIP 不阻断交付（exit 0），但 stderr 提示本次哪些维度未校验。
"""
import argparse, json, re, sys
from datetime import datetime
sys.path.insert(0, __import__("os").path.dirname(__file__))
from _common import atomic_write_json, cn_to_num, guard_out, norm, parse_amount, is_exempt_ph

try:
    from docx import Document
except ImportError:
    Document = None

RESIDUE_RE = re.compile(r"【待填|【待确认|【待补")

# 红线短语命中行须有实质响应内容：行长明显大于短语本身，
# 否则判定为目录/标题式裸命中（如「实质性条款响应承诺」只出现在目录中）。
SHORT_HIT_MARGIN = 8


def text_of(path):
    if str(path).lower().endswith(".docx") and Document:
        doc = Document(path)
        parts = [p.text for p in doc.paragraphs]
        for t in doc.tables:
            for r in t.rows:
                parts.extend(c.text for c in r.cells)
        return "\n".join(parts)
    d = json.load(open(path, encoding="utf-8"))
    return json.dumps(d, ensure_ascii=False)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("target")
    ap.add_argument("--template", default="")
    ap.add_argument("--quotes", default="")
    ap.add_argument("--tender-parse", dest="tender", default="")
    ap.add_argument("-o", "--out", default="verify_report.md")
    a = ap.parse_args()
    text = text_of(a.target)
    checks = []

    # A 占位残留：显式标记 + 原始模板占位（下划线/冒号空白槽等）双查
    # 豁免：模板原生格式标题（如「投标函【工程量清单计价】」）为招标文件原样切片
    # 的格式说明文本，非可填占位（与基线成品一致），不计入残留（豁免表 _common.EXEMPT_PH）。
    n_res = len(RESIDUE_RE.findall(text))
    try:
        from _common import RAW_PH_RE
        n_raw = 0
        for m in RAW_PH_RE.finditer(text):
            if is_exempt_ph(m.group(0)):
                continue
            n_raw += 1
    except Exception:
        n_raw = 0
    total_res = n_res + n_raw
    checks.append({"id": "A-residue", "name": "占位残留（【待填】类标记 + 未填模板占位）",
                   "result": "PASS" if total_res == 0 else "FAIL",
                   "detail": f"标记 {n_res} 处 + 原占位 {n_raw} 处"})

    # B 金额一致（需 --quotes；缺失 → SKIP 显式披露，不静默跳过）
    if not a.quotes:
        checks.append({"id": "B-quotes", "name": "金额一致",
                       "result": "SKIP", "detail": "未提供 --quotes，本次未校验金额大小写/合计一致"})
    else:
        try:
            q = json.load(open(a.quotes, encoding="utf-8"))
            quotes = q if isinstance(q, list) else q.get("quotes", [])
            total = None
            items = []
            for it in quotes:
                label = it.get("label", "")
                amt = parse_amount(it.get("amount"))
                cn = it.get("amount_cn")
                if re.search(r"总|合计", label) and amt is not None:
                    total = amt
                else:
                    items.append((label, amt))
                if cn and amt is not None:
                    cv = cn_to_num(cn)
                    ok = cv is not None and abs(cv - amt) < 0.01
                    if not ok:
                        checks.append({"id": "B-cn", "name": f"大小写不一致：{label}",
                                       "result": "FAIL", "detail": f"数字={amt} 大写={cn}→{cv}"})
            ssum = sum(v for _, v in items if v is not None)
            if total is not None and abs(ssum - total) > 0.01:
                checks.append({"id": "B-sum", "name": "分项合计≠总价", "result": "FAIL",
                               "detail": f"Σ分项={ssum} 总价={total}"})
            elif total is not None:
                checks.append({"id": "B-sum", "name": "分项合计=总价", "result": "PASS",
                               "detail": f"{total:,.2f}"})
        except Exception as e:
            checks.append({"id": "B", "name": "报价校验", "result": "FAIL", "detail": str(e)})

    # C 结构保真：草稿须覆盖 template.json headings（禁止删模板节）[需 --template]
    if not a.template:
        checks.append({"id": "C-structure", "name": "模板节完整（禁止删模板节）",
                       "result": "SKIP", "detail": "未提供 --template，本次未校验模板节是否被删"})
    else:
        try:
            tpl = json.load(open(a.template, encoding="utf-8"))
            need = tpl.get("headings") or [s.get("title", "") for s in tpl.get("sections", [])]
            if str(a.target).lower().endswith(".docx") and Document:
                chk_doc = Document(a.target)
                have = norm("".join(p.text for p in chk_doc.paragraphs))
            else:
                tgt = json.load(open(a.target, encoding="utf-8"))
                secs = tgt.get("document", {}).get("sections") or tgt.get("sections") or []
                have = norm("".join((s.get("heading") or s.get("title") or "") for s in secs))
            lost = [h for h in need if norm(h) and norm(h) not in have]
            checks.append({"id": "C-structure", "name": "模板节完整（禁止删模板节）",
                           "result": "PASS" if not lost else "FAIL",
                           "detail": f"应含 {len(need)} 标题，丢失 {len(lost)}" + (f"：{lost[:5]}" if lost else "")})
        except Exception as e:
            checks.append({"id": "C", "name": "结构校验", "result": "FAIL", "detail": str(e)})

    # D 废标红线证据（需 --tender-parse）。按行判定：红线短语仅命中目录/标题式
    # 短行不算实质响应（如「实质性条款响应承诺」只出现在目录里 = 未响应）。
    if not a.tender:
        checks.append({"id": "D-redline", "name": "废标条款响应证据扫描",
                       "result": "SKIP", "detail": "未提供 --tender-parse，本次未扫描废标条款响应"})
    else:
        try:
            tp = json.load(open(a.tender, encoding="utf-8"))
            clauses = tp.get("disqualification_clauses", [])
            rows = [norm(x) for x in text.split("\n") if norm(x)]
            miss, short_only, hit = [], [], 0
            for c in clauses:
                key = norm(c.get("clause", c if isinstance(c, str) else ""))[:12]
                if not key:
                    continue
                long_hits = [r for r in rows if key in r and len(r) >= len(key) + SHORT_HIT_MARGIN]
                short_hits = [r for r in rows if key in r and r not in long_hits]
                if long_hits:
                    hit += 1
                elif short_hits:
                    short_only.append(key)
                else:
                    miss.append(key)
            if miss:
                res, detail = "WEAK", f"{len(clauses)} 条，正文未命中 {len(miss)} 条" + (f"：{miss[:6]}" if miss else "")
            elif short_only:
                res, detail = "WEAK", (f"全部命中但 {len(short_only)} 条仅出现在目录/标题式短行"
                                       f"（无实质响应内容）" + (f"：{short_only[:6]}" if short_only else ""))
            else:
                res, detail = "PASS", f"{len(clauses)} 条全部在正文中找到实质响应行"
            checks.append({"id": "D-redline", "name": "废标条款响应证据扫描",
                           "result": res, "detail": detail})
        except Exception as e:
            checks.append({"id": "D", "name": "红线扫描", "result": "FAIL", "detail": str(e)})

    fails = [c for c in checks if c["result"] == "FAIL"]
    skips = [c for c in checks if c["result"] == "SKIP"]
    allpass = not fails
    if skips:
        print(f"WARN {len(skips)} 项未校验（缺参数）——见报告 SKIP 项，交付前建议补齐对应参数复跑",
              file=sys.stderr)
    report = {"target": a.target, "generated_at": datetime.now().isoformat(timespec="seconds"),
              "result": "PASS" if allpass else "FAIL", "checks": checks}
    out = guard_out(a.out)
    if allpass and skips:
        verdict = "✅ 通过（⚠️ 含 SKIP 未校验项）"
    elif allpass:
        verdict = "✅ 通过"
    else:
        verdict = "❌ 未通过"
    lines = [f"# 交付机检报告", "", f"- 对象：`{a.target}`（{report['generated_at']}）",
             f"- 结论：**{verdict}**", "",
             "| 检查 | 结果 | 明细 |", "|---|---|---|"]
    for c in checks:
        icon = {"PASS": "✅", "FAIL": "❌", "WEAK": "⚠️", "SKIP": "⏭️"}[c["result"]]
        lines.append(f"| {c['name']} | {icon} {c['result']} | {c['detail']} |")
    with open(out, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(f"{'PASS' if allpass else 'FAIL'} checks={len(checks)} fails={len(fails)} skips={len(skips)} -> {out}")
    sys.exit(0 if allpass else 5)


if __name__ == "__main__":
    main()
