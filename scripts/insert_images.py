#!/usr/bin/env python3
"""insert_images.py — 证照图片插入（环节二附）

从「成交响应文件 docx」中按小节抽取证照图片，插入到投标草稿对应小节标题段之后
（图片段居中、宽 6 英寸），避免用户手工扫描/贴图。

CLI:
  python insert_images.py <源成交响应.docx> <目标草稿.docx> -o 输出.docx \
      --map "营业执照:1,腐蚀控制资质:3,安取:1,质量体系:1,HSE:2"
    --map  小节标题关键词:该小节取前 n 张（逗号分隔；顺序=插入顺序）
退出码：0 完成；2 输入/依赖错误。目标 docx 被 Word/WPS 占用（PermissionError）
时自动另存 `_含图版.docx` 并提示，不覆盖原文件。

要点（2026-09 实践）：
- 映射建立：遍历源文档 body，遇「标题段落」切换当前小节；统计小节内全部 a:blip
  的 r:embed → {小节: [rid,...]}（标题段只在段落中匹配，表格误命中不切小节）。
- 图片抽取：rid → part blob 写盘（保留 partname 扩展名），每小节取前 n 张。
- 插入：python-docx 无 insert_after → 建空段 `insert_paragraph_before` 语义反用，
  以 `anchor._p.addnext(new_p)` 在锚点标题段后插入居中图段。
- 插入后必须 verify.py --template/--quotes/--tender-parse 复跑。
"""
import argparse, os, re, sys
sys.path.insert(0, os.path.dirname(__file__))
from _common import guard_out, norm

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches
    from docx.text.paragraph import Paragraph
except ImportError:
    print("ERROR: python-docx 未安装", file=sys.stderr)
    sys.exit(2)

IMG_WIDTH_IN = 6
MAX_TITLE_LEN = 40  # 标题段判定：含关键词且段长不超过该值


def parse_map(spec):
    """'营业执照:1,腐蚀控制资质:3' -> {'营业执照': 1, '腐蚀控制资质': 3}"""
    out = {}
    for part in (spec or "").split(","):
        part = part.strip()
        if not part:
            continue
        if ":" in part:
            k, _, n = part.rpartition(":")
        else:
            k, n = part, "1"
        if k:
            out[k.strip()] = max(int(n or 1), 1)
    return out


def collect_by_section(src_path, keys):
    """扫描源 docx：按小节标题统计 (key -> [rid,...])，保持出现顺序。"""
    doc = Document(src_path)
    mapping = {k: [] for k in keys}
    cur = None
    n_title_paras = 0
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            p = Paragraph(child, doc)
            t = norm(p.text).strip()
            if t and len(t) <= MAX_TITLE_LEN:
                for k in keys:
                    if norm(k) in t:
                        cur = k
                        n_title_paras += 1
                        break
            _collect_blips(p._p, mapping, cur)
        elif child.tag == qn("w:tbl"):
            for tc in child.iter(qn("w:tc")):
                for p_el in tc.iter(qn("w:p")):
                    _collect_blips(p_el, mapping, cur)
    return mapping, n_title_paras


def _collect_blips(p_el, mapping, cur):
    if cur is None:
        return
    for blip in p_el.iter(qn("a:blip")):
        rid = blip.get(qn("r:embed"))
        if rid:
            mapping[cur].append(rid)


def extract_blob(doc_part, rid, img_dir, key, idx):
    """rid → blob 写盘；返回图片绝对路径或 None。"""
    try:
        part = doc_part.related_parts[rid]
    except KeyError:
        return None
    blob = part.blob
    if not blob:
        return None
    ext = os.path.splitext(part.partname)[1].lower() or ".img"
    if ext not in (".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".emf", ".wmf"):
        ext = ".img"
    path = os.path.join(img_dir, f"{key}_{idx}{ext}")
    with open(path, "wb") as f:
        f.write(blob)
    return path


def insert_after_heading(doc, key, img_path):
    """在目标草稿中首个含 key 的标题段之后插入居中图段。返回 True/False。"""
    for p in doc.paragraphs:
        if norm(p.text).strip() and norm(key) in norm(p.text):
            new_p_el = OxmlElement("w:p")
            p._p.addnext(new_p_el)
            para = Paragraph(new_p_el, doc)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.add_run().add_picture(img_path, width=Inches(IMG_WIDTH_IN))
            return True
    return False


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("src", help="源成交响应文件 docx（抽图来源）")
    ap.add_argument("target", help="目标投标草稿 docx（插入对象）")
    ap.add_argument("-o", "--out", default="", help="输出路径（默认：目标_含图版.docx）")
    ap.add_argument("--map", required=True,
                    help="小节标题关键词:n 逗号分隔，如 '营业执照:1,腐蚀控制资质:3'")
    ap.add_argument("--img-dir", default="", help="抽图临时目录（默认：输出同目录 _assets/）")
    a = ap.parse_args()

    keys = parse_map(a.map)
    if not keys:
        print("ERROR: --map 未解析到任何小节", file=sys.stderr)
        sys.exit(2)
    mapping, n_title = collect_by_section(a.src, list(keys))
    if n_title == 0:
        print("WARN 源文档中未命中任何小节标题（--map 关键词与实际标题不一致？）", file=sys.stderr)

    out = a.out or re.sub(r"\.docx$", "", a.target) + "_含图版.docx"
    out_p = guard_out(out)
    img_dir = a.img_dir or os.path.join(os.path.dirname(out_p), "_assets")
    img_dir = guard_out(img_dir)
    os.makedirs(img_dir, exist_ok=True)

    src_doc = Document(a.src)
    dst_doc = Document(a.target)
    placed, missing_anchor, total_img = 0, [], 0
    for key in keys:
        n = mapping.get(key, []) and min(len(mapping[key]), keys[key]) or 0
        for i in range(n):
            img_path = extract_blob(src_doc.part, mapping[key][i], img_dir, key, i + 1)
            if not img_path:
                print(f"WARN {key} 第 {i+1} 张 rid 解析失败", file=sys.stderr)
                continue
            if insert_after_heading(dst_doc, key, img_path):
                placed += 1
                total_img += 1
            else:
                missing_anchor.append(key)
                break
    try:
        dst_doc.save(out_p)
    except PermissionError:
        alt = re.sub(r"\.docx$", "", out_p) + "_含图版.docx"
        try:
            dst_doc.save(alt)
            print(f"WARN 原输出被占用（Word 打开？），已另存：{alt}", file=sys.stderr)
            out_p = alt
        except PermissionError:
            print("ERROR: 输出被占用且另存失败，请关闭占用程序后重试", file=sys.stderr)
            sys.exit(2)
    print(f"OK 小节标题={n_title} 已插入={placed} 张={total_img} -> {out_p}")
    if missing_anchor:
        print(f"WARN 目标草稿中未找到锚点标题（未插入）" + (f"：{set(missing_anchor)}" if missing_anchor else ""),
              file=sys.stderr)


if __name__ == "__main__":
    main()
