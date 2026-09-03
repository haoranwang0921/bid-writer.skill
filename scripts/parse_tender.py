#!/usr/bin/env python3
"""parse_tender.py — 招标文件结构化解析（缺口 1，v0.2）

把招标/采购文件 docx 解析为结构化常量：
  * 常量：招标编号/项目名称/采购人/开标时间/有效期/工期/保证金
  * 评分项：评分表（表头含 分值/评分/评议）逐行提取（序号+评审项目+分值）+
            正文"分值构成"句（价格/商务/技术 权重）
  * 红线：评审/否决条款中含 否决/废标/不予受理 等的语句
  * 时间节点：带时间语义上下文的日期（过滤业绩年限假阳性）

输出：
  -o  tender.json       结构化解析结果
  -k  项目常量.json     轻量项目常量

退出码：0 成功；2 文件错误；3 解析为空（停下让用户核对）。
"""
import argparse
import json
import re
import sys
import os
from datetime import datetime

sys.path.insert(0, os.path.dirname(__file__))
from _common import atomic_write_json, guard_out, norm

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx 未安装", file=sys.stderr)
    sys.exit(2)

CN = "零一二三四五六七八九十百千"
CHAPTER_RE = re.compile(r"^\s*第\s*[" + CN + r"0-9１-９]+\s*[章节部分篇卷]")
NUM_LV2_RE = re.compile(r"^\s*\d{1,2}(?:[.．]\d{1,2})+\s*\S")

SCORE_ANCHOR = re.compile(r"评标办法|评分办法|评审办法|评分标准|评标细则|评审细则|分值构成")
REDLINE_ANCHOR = re.compile(r"否决|废标|否决其投标|不予受理|拒绝其投标|按无效投标|作废|取消其投标资格")
TIME_CTX_RE = re.compile(r"递交|截止|开标|获取|答疑|澄清|踏勘|公示|中标|签订|响应文件递交|保证金到账|退还")
COMPANY_RE = re.compile(r"公司|集团|局\b|中心|研究院|有限公司|有限责任公司|股份有限公司")
MONEY_RE = re.compile(r"([0-9][0-9,，]*)\s*(万|万元|元|亿元)")

SCORE_TABLE_HEAD = re.compile(r"分值|评分|评议|评分项|评审内容|评定标准")


def iter_texts(doc):
    """文档序 yield (kind, index, text)。kind: p/tbl。"""
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    body = doc.element.body
    idx = 0
    for ch in body:
        tag = ch.tag.split('}')[-1]
        if tag == 'p':
            yield ('p', idx, Paragraph(ch, doc).text)
            idx += 1
        elif tag == 'tbl':
            tb = Table(ch, doc)
            for r, row in enumerate(tb.rows):
                for c, cell in enumerate(row.cells):
                    t = cell.text
                    if t and t.strip():
                        yield ('tbl', (idx, r, c), t)
            idx += 1


def find_after(text, anchor, span_before=4, span_after=160):
    """找 anchor 之后的取值窗口；无则 None。"""
    m = re.search(anchor, text)
    if not m:
        return None
    start = max(0, m.start() - span_before)
    return text[start:m.start() + span_after]


def first_company_in(text):
    """在文本窗口里找第一个像公司名的片段（以公司/集团等收尾、≤40字）。"""
    if not text:
        return None
    for m in re.finditer(r"([\u4e00-\u9fa5A-Za-z0-9（）()]{3,40}?)(公司|集团|\S*局|中心|研究院)", text):
        cand = (m.group(1) + m.group(2)).strip("：: ：")
        if len(cand) >= 4 and COMPANY_RE.search(cand):
            return cand
    return None


def extract_constants(doc, joined):
    c = {}
    # 项目编号 / 招标编号
    for anchor in (r"项目编号[:：]", r"招标编号[:：]"):
        m = re.search(anchor + r"\s*([A-Za-z0-9][A-Za-z0-9\-_]{3,})", joined)
        if m:
            c["招标编号"] = m.group(1)
            break
    # 项目名称
    m = re.search(r"项目名称[:：]?\s*([^\n]{4,60})", joined[:3000])
    if m:
        c["项目名称"] = m.group(1).strip()
    # 采购人/招标人：优先匹配独立行的 `采购人：` / `招标人：`（跨行+公司名）
    for anchor_k, key in ((r"采购人", "采购人"), (r"招标人(?!名称)", "招标人"),
                          (r"发包人", "发包人")):
        win = find_after(joined, anchor_k, span_after=220)
        name = first_company_in(win)
        if name and key == "招标人" and "说明" in name[:4]:
            continue  # 命中的是"说明与XX分公司"残片，跳过
        if name:
            c[key] = name
            break
    if "采购人" not in c and "招标人" not in c:
        for p in doc.paragraphs[:120]:
            t = p.text.strip()
            m = re.search(r"^(?:采购人|招标人|招标人名称)\s*[:：]\s*([^\n]{4,60})$", t)
            if m:
                nm = m.group(1).strip()
                if COMPANY_RE.search(nm) or len(nm) >= 6:
                    c.setdefault("采购人", nm)
                    break
    # 开标 / 递交截止
    m = re.search(r"(?:开标[^。]{0,6}|递交[^。]{0,10}截止|响应文件[^。]{0,8}截止)[^0-9]{0,12}?"
                  r"([0-9]{4}\s*年\s*[0-9]{1,2}\s*月\s*[0-9]{1,2}\s*日)", joined)
    if m:
        c["开标时间"] = m.group(1)
    # 有效期（报价/响应/投标）
    m = re.search(r"(?:报价|响应|投标)\s*有效期[^0-9]{0,10}?([0-9]{1,3})\s*日历日", joined)
    if m:
        c["投标有效期_日历日"] = int(m.group(1))
    m = re.search(r"(?:报价|响应|投标)\s*有效期[^0-9]{0,10}?([0-9]{1,3})\s*[天日]", joined)
    if m:
        c.setdefault("投标有效期_日历日", int(m.group(1)))
    # 工期/服务期/供货期（仅在明确是"工期：N日历日/天"语义时取，排除合同条款噪声）
    for kw, key in ((r"工期", "工期_日历日"), (r"服务期", "服务期_日历日"), (r"供货期", "供货期_日历日")):
        m = re.search(kw + r"[^0-9]{0,8}?([0-9]{1,3})\s*(?:日历日|日历天|天)", joined)
        if m and not re.search(r"工期(?:调整报告|延长|延误|索赔|变更)", joined[max(0, m.start()-20):m.end()]):
            c[key] = int(m.group(1))
            break
    # 保证金
    m = re.search(r"保证金[^。]{0,40}?([0-9][0-9,，]{0,9})\s*(万元|元)", joined)
    if m:
        v = float(m.group(1).replace(",", "").replace("，", ""))
        c["保证金金额"] = round(v, 2)
        c["保证金单位"] = "万元" if m.group(2) == "万元" else "元"
    return c


def extract_evaluation(doc, joined):
    """评分项：评分表逐行 + 正文分值构成。"""
    items = []
    seen = set()

    # 1) 评分表：表头含 分值/评分/评议 → 行级提取（评审项目列=第2列，分值=末列）
    for idx, tb in enumerate(doc.tables):
        if not tb.rows:
            continue
        head = [c.text.strip() for c in tb.rows[0].cells]
        if not head or not any(SCORE_TABLE_HEAD.search(h) for h in head):
            continue
        if not any("分值" in h or "得分" in h or "评分" in h for h in head):
            continue
        for r in tb.rows[1:]:
            cells = [cc.text.strip() for cc in r.cells]
            if not cells:
                continue
            last = cells[-1]
            if not re.match(r"^\d{1,3}(\.\d)?$", last):  # 末列非纯数值则跳过（合计行等）
                continue
            score = float(last)
            name = ""
            # 评审项目列：找非序号、非分值、非合计 的短文本列
            for cc in cells[1:-1]:
                t = cc.replace("\n", "")
                if re.match(r"^\d{1,3}$", t) or t in ("合计", "") or len(t) < 1:
                    continue
                name = t
                break
            if not name:
                continue
            k = norm(name)
            if k in seen:
                continue
            seen.add(k)
            items.append({"no": cells[0] if re.match(r"^\d+$", cells[0]) else "",
                          "item": name[:80], "score": score,
                          "source": f"表{idx}"})

    # 2) 正文分值构成句（如 价格评审50分，商务评审15分，技术评审35分）
    m = re.search(r"(总分为100分|满分100|总分100).{0,60}?价格[^\n，。]{0,12}?([0-9]{1,3})"
                  r"[^，。]{0,12}?商务[^\n，。]{0,12}?([0-9]{1,3})"
                  r"[^，。]{0,12}?技术[^\n，。]{0,12}?([0-9]{1,3})", joined)
    if m:
        for k, v in (("价格评审", m.group(2)), ("商务评审", m.group(3)), ("技术评审", m.group(4))):
            kk = norm(k)
            if kk not in seen:
                seen.add(kk)
                items.append({"no": "", "item": k, "score": float(v), "source": "分值构成"})
    return items[:80]


def extract_redlines(doc):
    """废标红线：含否决/废标等关键词的语句（含表格文本）。"""
    out, seen = [], set()
    cur_ch = "(文档开头)"
    for item in iter_texts(doc):
        kind, key, text = item
        t = text.strip()
        if kind == 'p':
            if CHAPTER_RE.match(t) and len(t) <= 40:
                cur_ch = t
                continue
        if len(t) > 5 and REDLINE_ANCHOR.search(t):
            k = norm(t[:40])
            if k in seen:
                continue
            seen.add(k)
            out.append({"clause": t[:220], "chapter": cur_ch})
    return out[:80]


def extract_timeline(joined):
    """时间节点：日期上下文含时间语义词才保留（过滤 2023 业绩年限假阳性）。"""
    out, seen = [], set()
    for m in re.finditer(r"([0-9]{4})\s*年\s*([0-9]{1,2})\s*月\s*([0-9]{1,2})\s*日", joined):
        date = f"{m.group(1)}年{m.group(2)}月{m.group(3)}日"
        ctx = joined[max(0, m.start() - 40): m.start()].strip()
        ctx = ctx.replace("\n", " ")
        if not TIME_CTX_RE.search(ctx[-36:]):
            continue
        k = norm(date + ctx[-24:])
        if k in seen or date in seen:
            continue
        seen.add(k)
        out.append({"date": date, "ctx": ctx[-40:], "found_in": "正文"})
        if len(out) >= 12:
            break
    # 优先排序：2026（开标年前后）在前
    out.sort(key=lambda x: (x["date"] <= "2025", x["ctx"]))
    return out


def parse(doc):
    txts = [i[2] for i in iter_texts(doc) if i[0] == 'p']
    joined = "\n".join(t for t in txts if t)
    constants = extract_constants(doc, joined)
    evaluation = extract_evaluation(doc, joined)
    redlines = extract_redlines(doc)
    timeline = extract_timeline(joined)
    return constants, evaluation, redlines, timeline


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("tender")
    ap.add_argument("-o", "--out", default="tender.json")
    ap.add_argument("-k", "--const", default="")
    ap.add_argument("--with-timeline", action="store_true")
    a = ap.parse_args()
    try:
        doc = Document(a.tender)
    except Exception as e:
        print(f"ERROR: 无法打开 {a.tender}: {e}", file=sys.stderr)
        sys.exit(2)
    constants, ev, red, tl = parse(doc)
    if not constants and not ev and not red:
        print("MISS 未解析出有效常量/评分/红线 —— 请人工核对招标文件或调整解析锚点", file=sys.stderr)
        sys.exit(3)
    data = {
        "source": a.tender,
        "parsed_at": datetime.now().isoformat(timespec="seconds"),
        "constants": constants,
        "evaluation_items": ev,
        "disqualification_clauses": red,
        "timeline": tl if a.with_timeline else [],
        "stats": {"constants": len(constants), "evaluation": len(ev),
                  "redlines": len(red), "timeline": len(tl)},
    }
    out = guard_out(a.out)
    atomic_write_json(out, data)
    if a.const:
        kout = guard_out(a.const)
        atomic_write_json(kout, {"招标编号": constants.get("招标编号", ""),
                                 "项目名称": constants.get("项目名称", ""),
                                 "采购人名称": constants.get("采购人", constants.get("招标人", "")),
                                 "开标时间": constants.get("开标时间", ""),
                                 "投标有效期_日历日": constants.get("投标有效期_日历日", 90),
                                 "工期_日历日": constants.get("工期_日历日", 0)})
        print(f"OK 常量={len(constants)} 评分项={len(ev)} 废标红线={len(red)} 时间节点={len(tl)}")
        print(f"   tender.json -> {out}")
        print(f"   项目常量.json -> {kout}")
    else:
        print(f"OK 常量={len(constants)} 评分项={len(ev)} 废标红线={len(red)} 时间节点={len(tl)} -> {out}")
    if red:
        print("   红线样例：")
        for r in red[:3]:
            print(f"     [{r['chapter'][:12]}] {r['clause'][:60]}")


if __name__ == "__main__":
    main()