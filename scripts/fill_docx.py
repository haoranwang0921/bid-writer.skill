#!/usr/bin/env python3
"""fill_docx.py — 将确认值就地写回模板切片（局部修改，格式继承自源文件）

读取 template.docx 副本 + fills.json，把已确认值替换进对应占位；未确认占位
保持原样。输出即为投标文件底稿（后续叙述内容由 agent 在同文件上按模板小节
续写，禁止增删模板既有节）。

**格式保真策略（v0.3.2）**：
- 优先「单 run 内替换」：找到含占位文本的 run，直接 `r.text.replace(...)` →
  整 run 属性（rPr：下划线/字号/字体/底纹）天然保留，**这是 99% 场景**。
- 占位跨多 run（Word 自动拆字、合并字符）：把填入文字构造为新 run，
  **复制占位 run 的 rPr**，原占位 run 全部置空（保持段落长度不变，避免
  后续 fill_plan/find 索引错位）。
- 不重建文档，不重排段落，不重设样式。

CLI:
  python fill_docx.py <template.docx> <fills.json> -o 投标文件草稿_YYYYMMDD.docx
退出码：0 全部填入；3 仍有未填占位（交付前必须清零）。
"""
import argparse, json, re, sys, os
sys.path.insert(0, os.path.dirname(__file__))
from _common import guard_out

try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx 未安装", file=sys.stderr)
    sys.exit(2)

PH_RE = re.compile(r"_{2,}|＿{2,}|×{3,}|【[^】]{0,20}】|（\s*）|\(\s*\)|[：:][ \t\u3000]{2,}(?=[（(年元天日%]|$)")


def _copy_rpr(src_rpr):
    """深复制源 rPr（XML 节点级），返回独立副本。返回 None 当源不存在。"""
    if src_rpr is None:
        return None
    from copy import deepcopy
    return deepcopy(src_rpr)


def _replace_in_para(p, old, new):
    """run 级替换优先；跨 run 则把填入文字构为新 run（带占位 run 的 rPr）。

    返回 True 表示替换成功（无论单 run 还是跨 run）。
    """
    # 1) 单 run 内包含全部 old：直接改 r.text，rPr 自然保留
    for r in p.runs:
        if old in r.text:
            r.text = r.text.replace(old, new, 1)
            return True
    # 2) 跨 run：先确定 old 在段落里的起止 run
    full = p.text
    idx = full.find(old)
    if idx < 0:
        return False
    # 用段落元素的 run 列表按文本起止位置切分（python-docx 的 .runs 不准
    # ——这里直接遍历 XML w:r 子元素，按累计文本长度比对）。
    p_el = p._p
    r_elements = p_el.findall(qn("w:r"))
    if not r_elements:
        return False
    # 计算每个 run 的 [start, end) 文本位置
    pos = 0
    spans = []
    for r_el in r_elements:
        t_els = r_el.findall(qn("w:t"))
        r_text = "".join(t.text or "" for t in t_els)
        spans.append((r_el, pos, pos + len(r_text), r_text))
        pos += len(r_text)
    old_start, old_end = idx, idx + len(old)
    in_runs = [(el, s, e, t) for (el, s, e, t) in spans if not (e <= old_start or s >= old_end)]
    if not in_runs:
        return False
    # 取首个包含 old 的 run 作为「占位 run」——其 rPr 是下划线等格式的来源
    anchor_rpr = in_runs[0][0].find(qn("w:rPr"))
    # 替换为「new」：策略是
    # a) 把 in_runs 内首个 run 的文本改为「new」+ 紧贴 old_start 之前的尾部
    # b) 其他 in_runs 的 run 全部置空
    # 这样：填入文字落在带原始 rPr 的首个 run 上，原下划线格式继承
    first_el, first_s, first_e, first_text = in_runs[0]
    # 首 run 内 old 之前的部分 + new + old 之后的部分
    pre = first_text[: old_start - first_s]
    post = first_text[old_start - first_s + len(old):]
    # 重写首 run 的 w:t 为 pre+new（post 仍属于同一段，继续放在第二个新建 run 中
    # 保留原 rPr）。
    new_first_text = pre + new
    # 移除首 run 的所有 w:t，写入新文本
    for t in first_el.findall(qn("w:t")):
        first_el.remove(t)
    nt = OxmlElement("w:t")
    nt.text = new_first_text
    nt.set(qn("xml:space"), "preserve")
    first_el.append(nt)
    # 处理 post 段：放在首 run 之后的新 run，rPr 复制
    if post:
        new_r = OxmlElement("w:r")
        rpr_copy = _copy_rpr(anchor_rpr)
        if rpr_copy is not None:
            new_r.append(rpr_copy)
        nt2 = OxmlElement("w:t")
        nt2.text = post
        nt2.set(qn("xml:space"), "preserve")
        new_r.append(nt2)
        first_el.addnext(new_r)
    # 其余 in_runs（>=1）置空
    for el, _, _, _ in in_runs[1:]:
        for t in el.findall(qn("w:t")):
            el.remove(t)
    return True


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("template")
    ap.add_argument("fills")
    ap.add_argument("-o", "--out", default=None)
    a = ap.parse_args()
    doc = Document(a.template)
    fills = json.load(open(a.fills, encoding="utf-8"))
    fmap = fills.get("fills", fills) if isinstance(fills, dict) else {}
    done = miss = 0
    # 段落槽
    for pid, val in list(fmap.items()):
        m = re.match(r"^P(\d+)#(\d+)$", pid)
        if not m:
            continue
        pi, k = int(m.group(1)), int(m.group(2))
        if pi >= len(doc.paragraphs):
            miss += 1
            continue
        p = doc.paragraphs[pi]
        ms = list(PH_RE.finditer(p.text or ""))
        if k < len(ms):
            if _replace_in_para(p, ms[k].group(0), str(val)):
                done += 1
            else:
                miss += 1
    # 表格槽：建立 (ti,r,c)→cell 稳定映射
    tcoord = {}
    for ti, tb in enumerate(doc.tables):
        for r, row in enumerate(tb.rows):
            for c, cell in enumerate(row.cells):
                tcoord[(ti, r, c)] = cell
    for pid, val in list(fmap.items()):
        m = re.match(r"^T(\d+):R(\d+)C(\d+)$", pid)
        if not m:
            continue
        cell = tcoord.get((int(m.group(1)), int(m.group(2)), int(m.group(3))))
        if cell is None:
            miss += 1
            continue
        for p in cell.paragraphs:
            for ph in list(PH_RE.finditer(p.text or "")):
                if _replace_in_para(p, ph.group(0), str(val)):
                    done += 1
                    break
                break
    out = a.out or re.sub(r"\.docx$", "", a.template) + "_草稿.docx"
    out = guard_out(out)
    doc.save(out)
    residual = 0
    chk = Document(out)
    for p in chk.paragraphs:
        residual += len(PH_RE.findall(p.text or ""))
    for tb in chk.tables:
        for row in tb.rows:
            for cell in row.cells:
                residual += len(PH_RE.findall(cell.text or ""))
    print(f"OK -> {out} 填入={done} 定位失败={miss} 剩余占位={residual}")
    if residual:
        print("WARN 剩余占位须在提问确认后清零，否则不可交付（W2/W4）", file=sys.stderr)
        sys.exit(3)


if __name__ == "__main__":
    main()
