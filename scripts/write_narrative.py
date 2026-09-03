#!/usr/bin/env python3
"""write_narrative.py — 叙述章节撰写（缺口 2）

对投标/响应文件草稿中"要求投标人自拟"的叙述性小节（施工组织设计、技术方案、
质量/安全/进度/环保措施、项目管理机构说明等），用本地知识库检索佐证素材，
生成叙述文本并**就地追加到草稿对应小节末尾**；同时产出溯源 JSON。

叙述章节的识别基于 template.json 的 headings：
  * 命中叙述关键词（施工组织设计/技术方案/措施计划/方案/应急预案/组织管理机构…）
  * 或 --patterns 提供（对 seeker 型招标：施工方案/质量保证/安全措施…）

本脚本不自作文本，只做「检索 + 按小节拼接知识库引用块 + 写入草稿」；
真正语言组织由 agent 在产出文本基础上修改（数字/参数必须溯源）。

用法：
  python write_narrative.py <草稿.docx> <template.json> \
      [--out 草稿叙述版.docx] [--pattern 施工方案,质量保证,安全措施] \
      [--service http://127.0.0.1:8765] [--top-k 3] [--max-lines 6] [--kb json]

输出：
  -o       草稿叙述版.docx（在模板叙述小节后追加【知识库引用】块）
  --kb     叙述_溯源.json（heading → 引用片段列表）

退出码：0 成功；2 输入错误；3 草稿/模板无有效叙述小节可写（不视为失败，
        由调用方决定是否继续）；服务不可达按 warn 跳过（不阻断）。
"""
import argparse
import json
import os
import re
import sys
import urllib.error
import urllib.request
from datetime import datetime

sys.path.insert(0, os.path.dirname(__file__))
from _common import atomic_write_json, guard_out, norm

try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx 未安装", file=sys.stderr)
    sys.exit(2)

NARRATIVE_RE = re.compile(
    r"施工组织设计|施工方案|技术方案|供货和服务方案|质量保证|质量措施|安全措施|进度计划|"
    r"环保措施|环境措施|文明施工|应急预案|风险管理|组织管理机构|项目管理机构|"
    r"成品保护|保修|冬季|雨季|配合|协调|附表")
PATTERN_ANY = None


def query_service(service_url, query, top_k=3, library=None, min_score=0.35):
    """调用 /query 返回 results 列表；失败返回 []。"""
    payload = {"query": query, "top_k": top_k, "min_score": min_score}
    if library:
        payload["library"] = library
    data = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    req = urllib.request.Request(
        f"{service_url}/query", data=data,
        headers={"Content-Type": "application/json"})
    try:
        with urllib.request.urlopen(req, timeout=10) as resp:
            out = json.loads(resp.read().decode("utf-8"))
        return out.get("results", [])
    except (urllib.error.URLError, OSError, json.JSONDecodeError, ValueError):
        return []


def pick_narrative_headings(template_json, patterns=None):
    """从 template.json headings 挑叙述小节标题（按出现序去重）。"""
    heads = template_json.get("headings") or []
    out, seen = [], set()
    pat_re = None
    if patterns:
        pat_re = re.compile("|".join(re.escape(p) for p in patterns))
    for h in heads:
        if not h or len(h) > 60:
            continue
        hit = pat_re.search(h) if pat_re else NARRATIVE_RE.search(h)
        if hit:
            hk = norm(h)
            if hk in seen:
                continue  # 模板清单与格式件清单会重复出现同一标题，只写一次
            seen.add(hk)
            out.append(h)
    return out


def find_heading_para(doc, target):
    """在草稿中定位与目标标题 norm 相等的段落（限首段），返回 Paragraph。"""
    tn = norm(target)
    for p in doc.paragraphs:
        if norm(p.text) == tn:
            return p
    return None


def find_insert_point(body_parent, anchor_el):
    """返回锚点元素之后可插入兄弟元素的位置。"""
    return anchor_el


def add_para_after(doc, anchor_el, text, style=None):
    """在 anchor_el 之后插入新段落（与其同级）。"""
    from copy import deepcopy
    from docx.text.paragraph import Paragraph
    new_p = OxmlElement("w:p")
    anchor_el.addnext(new_p)
    p = Paragraph(new_p, doc)
    if style:
        try:
            p.style = style
        except Exception:
            pass
    run = p.add_run(text)
    run.font.size = None
    return p


def main():
    global PATTERN_ANY
    ap = argparse.ArgumentParser()
    ap.add_argument("draft", help="投标草稿 docx（fill_docx 产物）")
    ap.add_argument("template_json", help="template_v2.json（提取元数据）")
    ap.add_argument("--out", default=None, help="输出草稿（默认 输入_叙述版.docx）")
    ap.add_argument("--patterns", default="", help="自拟叙述小节关键词（逗号分隔）")
    ap.add_argument("--service", default="http://127.0.0.1:8765")
    ap.add_argument("--top-k", type=int, default=3)
    ap.add_argument("--max-lines", type=int, default=6, help="每个小节最多引用片段数")
    ap.add_argument("--min-score", type=float, default=0.35)
    ap.add_argument("--kb", default="", help="同时输出 叙述_溯源.json（引用档案）")
    a = ap.parse_args()

    try:
        doc = Document(a.draft)
        tpl = json.load(open(a.template_json, encoding="utf-8"))
    except Exception as e:
        print(f"ERROR: {e}", file=sys.stderr)
        sys.exit(2)

    heads = pick_narrative_headings(tpl, 
                                    [p.strip() for p in a.patterns.split(",") if p.strip()] or None)
    if not heads:
        print("MISS 未识别到叙述性小节（template.json headings 无命中，或 --patterns 无效）",
              file=sys.stderr)
        print("  提示：若模板本身无叙述章节则跳过本环节，进入 verify/diff", file=sys.stderr)
        sys.exit(3)

    archive, written, nohit = [], 0, 0

    for h in heads:
        # 1) 定位草稿中该小节
        anchor = find_heading_para(doc, h)
        if anchor is None:
            nohit += 1
            continue
        # 2) 检索佐证素材
        results = query_service(a.service, h, top_k=a.top_k, min_score=a.min_score)
        if not results:
            nohit += 1
            continue
        # 3) 拼接引用块文本：先【知识库引用 N】注释行 + 片段正文（可读、可溯源）
        lines = []
        for i, r in enumerate(results[:a.max_lines], 1):
            meta = r.get("metadata") or {}
            src = meta.get("source_path", "未知来源")
            heading = meta.get("heading_path", "")
            score = r.get("score", 0)
            lines.append(f"【知识库引用 {i}】相似度{score:.2f} | 来源:{src}")
            if heading:
                lines.append(f"章节:{heading}")
            lines.append(r.get("text", "")[:400])
            lines.append("")
            archive.append({"heading": h, "score": score, "source": src,
                            "heading_path": heading, "text": r.get("text", "")[:400]})
        block = "\n".join(lines).rstrip()
        # 4) 写入草稿：标题下方插入引用块（游标逐段追加，保持自然顺序）
        lines = ["[参考素材—请据此起稿，可删除本行]"] + block.split("\n")
        cursor = anchor._p
        for text in lines:
            new_p = OxmlElement("w:p")
            cursor.addnext(new_p)
            p = type(anchor)(new_p, doc)
            p.add_run(text)
            cursor = new_p
        written += 1

    out = a.out or re.sub(r"\.docx$", "", a.draft) + "_叙述版.docx"
    out_p = guard_out(out)
    try:
        doc.save(out_p)
    except PermissionError:
        print(f"WARN 输出文件被占用（Word 打开？）无法保存：{out_p}", file=sys.stderr)
        sys.exit(2)
    print(f"OK 叙述小节={len(heads)} 已写入={written} 检索无命中={nohit} -> {out_p}")
    if a.kb:
        kb_p = guard_out(a.kb)
        atomic_write_json(kb_p, {
            "draft": a.draft, "generated_at": datetime.now().isoformat(timespec="seconds"),
            "service": a.service, "total_heads": len(heads), "written": written,
            "retrieved": archive})
        print(f"   溯源 -> {kb_p}")
    if written == 0:
        print("WARN 所有叙述小节均未写入（服务不可达或检索无命中）——请检查知识库服务",
              file=sys.stderr)
        sys.exit(3)


if __name__ == "__main__":
    main()