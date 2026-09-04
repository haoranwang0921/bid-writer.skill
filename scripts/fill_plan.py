#!/usr/bin/env python3
"""fill_plan.py v2 — 置信度分级填空（对切片 template.docx 操作）

槽位 = template.docx 中的占位（段落空白槽/下划线/【】 与表格占位单元格）。
CLI:
  python fill_plan.py scan <template.docx> [--sources a.json,b.json] \
      -o fill_plan.json [--questions questions.json]
  python fill_plan.py apply <template.docx> <fill_plan.json> \
      [--smart smart_fill.json] [--answers answers.json] -o fills.json

置信度：
  high=画像键名与槽位 label 归一化后【精确一致】（扁平键末段或整键）且值唯一 → auto；
  medium=精确键多值冲突 / 仅子串（模糊）命中 → ask（语义复核，防串值）；
  low=无任何来源 → ask（禁止臆测）。
fills.json：{"fills": {"P12#0": "值", "T3:R1C2": "值"}}，并记录每个值来自
deterministic_auto / model_evidence / human_answer。优先级：人工 > 已验证模型 > 机械 high。
"""
import argparse, json, re, sys, os
from datetime import datetime
sys.path.insert(0, os.path.dirname(__file__))
from _common import atomic_write_json, guard_out, norm, is_exempt_ph

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx 未安装", file=sys.stderr)
    sys.exit(2)

PH_RE = re.compile(r"_{2,}|＿{2,}|×{3,}|【[^】]{0,20}】|（\s*）|\(\s*\)|[：:][ \t\u3000]{2,}(?=[（(年元天日%]|$)")
UNFILLABLE_RE = re.compile(r"待确认|待补|见扫描件|按招标格式")  # 画像中的占位值不得视为高置信度


def flatten(obj, prefix=""):
    out = {}
    if isinstance(obj, dict):
        for k, v in obj.items():
            key = f"{prefix}.{k}" if prefix else str(k)
            out.update(flatten(v, key) if isinstance(v, (dict, list)) else {key: v})
    elif isinstance(obj, list):
        for i, v in enumerate(obj):
            out.update(flatten(v, f"{prefix}[{i}]"))
    else:
        out[prefix or "value"] = obj
    return out


def build_slots(doc):
    slots = []
    for pi, p in enumerate(doc.paragraphs):
        t = p.text
        for k, m in enumerate(PH_RE.finditer(t or "")):
            if is_exempt_ph(m.group(0)):
                continue  # 模板原生格式标题（如【工程量清单计价】），非可填占位
            lead = norm(t[max(0, m.start() - 18):m.start()])[-16:]
            slots.append({"id": f"P{pi}#{k}", "type": "paragraph",
                          "label": lead or norm(t)[:20], "current": m.group(0),
                          "context": norm(t)[:70]})
    for ti, tb in enumerate(doc.tables):
        header = [c.text.strip() for c in tb.rows[0].cells] if len(tb.rows) else []
        seen_row = set()
        for r, row in enumerate(tb.rows):
            for c, cell in enumerate(row.cells):
                if (r, id(cell._tc)) in seen_row:
                    continue
                seen_row.add((r, id(cell._tc)))
                txt = cell.text.strip()
                if txt and PH_RE.search(txt):
                    exempt_cell = all(is_exempt_ph(m.group(0)) for m in PH_RE.finditer(txt))
                    if exempt_cell:
                        continue
                    lab = norm(header[c])[:24] if c < len(header) and header[c] else norm(txt)[:24]
                    slots.append({"id": f"T{ti}:R{r}C{c}", "type": "table_cell",
                                  "label": lab, "current": txt[:40]})
    return slots


def _leaf(key):
    """扁平键末段：profile.company.name → name（供精确匹配槽位 label）。"""
    return key.rsplit(".", 1)[-1]


def _exact_hit(key, lab):
    """精确命中判定：整键相等，或 label 以键结尾且前缀仅为短编号（（九）/1.2/一、/第X章）。
    编号前缀不改变键名语义，视为精确；非编号前缀（如「项目」+「名称」串值场景）一律
    不算精确，落入 fuzzy → medium，杜绝泛词键错填。"""
    if key == lab:
        return True
    if len(lab) <= len(key) or not lab.endswith(key):
        return False
    prefix = lab[: len(lab) - len(key)]
    num_chars = set("（）()一二三四五六七八九十百零壹贰叁肆伍陆柒捌玖拾0123456789.、．第章节条款")
    return bool(prefix) and len(prefix) <= 8 and all(ch in num_chars for ch in prefix)


def match_sources(slots, sources):
    # exact_idx 同时索引「整键」与「末段键」，归一化后与槽位 label 精确比较
    idx = {}
    for path in sources:
        try:
            data = json.load(open(path, encoding="utf-8"))
        except Exception as e:
            print(f"WARN 资料读取失败 {path}: {e}", file=sys.stderr)
            continue
        for k, v in flatten(data).items():
            if v is None or isinstance(v, bool) or str(v).strip() == "":
                continue
            if UNFILLABLE_RE.search(str(v)) or str(k).lower().endswith(("profile_id", "review_status")):
                continue  # 画像占位值/元字段不作为可填来源（自然落入提问档）
            rec = {"value": str(v), "source": os.path.basename(path)}
            full, leaf = norm(k), norm(_leaf(k))
            if len(full) >= 2:
                idx.setdefault(full, []).append(rec)
            if leaf and leaf != full and len(leaf) >= 2:
                idx.setdefault(leaf, []).append(rec)
    for s in slots:
        # label 常为「键名 + 冒号 + 空白」（项目名称：____），匹配比较时剥除尾冒号
        lab = norm(s["label"]).rstrip(":").strip()
        cands, seen = [], []
        if len(lab) >= 2:
            # 第一优先：精确键命中（整键相等，或 label 以键结尾且前缀≤6 短编号）
            for key, vals in idx.items():
                if _exact_hit(key, lab):
                    for v in vals:
                        if v["value"] not in seen:
                            seen.append(v["value"])
                            cands.append(dict(v, exact=True))
            # 其次：子串命中仅作候选提示（exact=False），永不直接判定 high
            if not cands:
                for key, vals in idx.items():
                    if len(key) >= 2 and key != lab and (lab in key or key in lab):
                        for v in vals:
                            if v["value"] not in seen:
                                seen.append(v["value"])
                                cands.append(dict(v, exact=False))
        s["candidates"] = cands[:4]
        exact_vals = {c["value"] for c in cands if c.get("exact")}
        has_fuzzy = any(not c.get("exact") for c in cands)
        if len(exact_vals) == 1 and not has_fuzzy:
            s["confidence"] = "high"
        elif cands:
            s["confidence"] = "medium"
        else:
            s["confidence"] = "low"
        s["status"] = "auto" if s["confidence"] == "high" else "ask"
    return slots


def cmd_scan(a):
    doc = Document(a.template)
    sources = [p.strip() for p in a.sources.split(",") if p.strip()] if a.sources else []
    slots = match_sources(build_slots(doc), sources)
    stats = {"total": len(slots), "high": 0, "medium": 0, "low": 0}
    for s in slots:
        stats[s["confidence"]] += 1
    source_catalog = [
        {"id": f"json-{i + 1}", "kind": "json", "location": os.path.abspath(path),
         "source_type": "structured", "access": "read_only"}
        for i, path in enumerate(sources)
    ]
    atomic_write_json(guard_out(a.out), {"meta": {"template": a.template,
                     "generated_at": datetime.now().isoformat(timespec="seconds"), "stats": stats,
                     "source_catalog": source_catalog},
                     "slots": slots})
    if a.questions:
        qs = [{"id": s["id"], "label": s["label"], "current": s["current"],
               "context": s.get("context", ""), "confidence": s["confidence"],
               "options": [f"{c['value']}（依据：{c['source']}）" for c in s["candidates"]]
                          + (["自行填写"] if s["confidence"] == "low" else []),
               "basis": "来源冲突需裁决" if s["confidence"] == "medium" else "无来源，禁止臆测"}
              for s in slots if s["status"] == "ask"]
        atomic_write_json(guard_out(a.questions), {"questions": qs})
    asks = sum(1 for s in slots if s["status"] == "ask")
    print(f"OK slots={stats['total']} high={stats['high']} medium={stats['medium']} "
          f"low={stats['low']} 待确认={asks} -> {a.out}")


def load_validated_smart(path):
    """读取 smart_fill.validate 产物，并再次核对 auto_fills 与决策审计一致。"""
    if not path:
        return {}, {}
    data = json.load(open(path, encoding="utf-8"))
    meta = data.get("meta", {})
    if meta.get("protocol") != "bid-writer-smart-fill/v1" or meta.get("validated") is not True:
        raise ValueError("--smart 不是经 smart_fill.py validate 生成的受控决策文件")
    decisions = {d.get("slot_id"): d for d in data.get("decisions", [])
                 if isinstance(d, dict) and d.get("slot_id")}
    auto, audit = {}, {}
    for sid, value in data.get("auto_fills", {}).items():
        d = decisions.get(sid, {})
        check = d.get("validation", {})
        if (check.get("auto_eligible") is not True
                or d.get("effective_confidence") != "high"
                or str(d.get("selected_value", "")) != str(value)):
            raise ValueError(f"--smart 槽位 {sid} 的 auto_fills 与验证记录不一致")
        auto[sid] = str(value)
        audit[sid] = {
            "origin": "model_evidence",
            "canonical_field": d.get("canonical_field", ""),
            "query": d.get("query", ""),
            "reason": d.get("reason", ""),
            "evidence": d.get("evidence", []),
        }
    return auto, audit


def cmd_apply(a):
    doc = Document(a.template)
    plan = json.load(open(a.plan, encoding="utf-8"))
    answers = {}
    if a.answers:
        raw = json.load(open(a.answers, encoding="utf-8"))
        answers = raw.get("answers", raw)
    smart_fills, smart_audit = load_validated_smart(a.smart)
    known_ids = {s["id"] for s in plan["slots"]}
    unknown_smart = sorted(set(smart_fills) - known_ids)
    if unknown_smart:
        raise ValueError(f"--smart 含未知槽位：{','.join(unknown_smart)}")
    fills, audit, unconf, unconfirmed_ids = {}, {}, 0, []
    for s in plan["slots"]:
        if s["id"] in answers:
            fills[s["id"]] = str(answers[s["id"]])
            audit[s["id"]] = {"origin": "human_answer"}
        elif s["id"] in smart_fills:
            fills[s["id"]] = smart_fills[s["id"]]
            audit[s["id"]] = smart_audit[s["id"]]
        elif s["status"] == "auto":
            fills[s["id"]] = s["candidates"][0]["value"]
            audit[s["id"]] = {
                "origin": "deterministic_auto",
                "evidence": s.get("candidates", [])[:1],
            }
        else:
            unconf += 1
            unconfirmed_ids.append(s["id"])
    atomic_write_json(guard_out(a.out), {"fills": fills, "audit": audit,
                     "unconfirmed": unconf, "unconfirmed_ids": unconfirmed_ids,
                     "generated_at": datetime.now().isoformat(timespec="seconds")})
    print(f"OK fills={len(fills)} 未确认={unconf} -> {a.out}")
    if unconf:
        print("WARN 未确认槽位将保持占位原样，交付前必须清零", file=sys.stderr)


def main():
    ap = argparse.ArgumentParser()
    sub = ap.add_subparsers(dest="cmd", required=True)
    s1 = sub.add_parser("scan")
    s1.add_argument("template")
    s1.add_argument("--sources", default="")
    s1.add_argument("-o", "--out", default="fill_plan.json")
    s1.add_argument("--questions", default="")
    s1.set_defaults(fn=cmd_scan)
    s2 = sub.add_parser("apply")
    s2.add_argument("template")
    s2.add_argument("plan")
    s2.add_argument("--smart", default="", help="smart_fill.py validate 产物")
    s2.add_argument("--answers", default="")
    s2.add_argument("-o", "--out", default="fills.json")
    s2.set_defaults(fn=cmd_apply)
    a = ap.parse_args()
    try:
        a.fn(a)
    except Exception as e:
        print(f"ERROR: {type(e).__name__}: {e}", file=sys.stderr)
        sys.exit(2)


if __name__ == "__main__":
    main()
