# -*- coding: utf-8 -*-
"""fill_docx 下划线保真回归测试（v0.3.2）。

锁定行为：填入文字必须继承占位 run 的下划线 rPr，无论单 run 还是跨 run 场景。
"""
import os
import tempfile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

import fill_docx as fd


def _add_underline(rPr):
    u = rPr.find(qn("w:u"))
    if u is None:
        u = etree.SubElement(rPr, qn("w:u"))
    u.set(qn("w:val"), "single")
    return rPr


def _mk_para(doc, runs_spec):
    """runs_spec = [(text, dict)]  None 或 {} = 不设格式。"""
    p = doc.add_paragraph()
    for text, props in runs_spec:
        r = p.add_run(text)
        if props:
            rPr = r._element.find(qn("w:rPr"))
            if rPr is None:
                rPr = OxmlElement("w:rPr")
                r._element.insert(0, rPr)
            if props.get("underline"):
                _add_underline(rPr)
    return p


def _has_underline(run):
    rPr = run._element.find(qn("w:rPr"))
    return rPr is not None and rPr.find(qn("w:u")) is not None


def _assert_underline_in(para, must_substr):
    for r in para.runs:
        if must_substr in r.text:
            assert _has_underline(r), f"段 '{para.text}' run「{r.text}」下划线丢失"
            return
    raise AssertionError(f"段 '{para.text}' 中找不到「{must_substr}」")


def _mk_template(tmp):
    path = os.path.join(tmp, "tpl.docx")
    doc = Document()
    _mk_para(doc, [("项目名称：____", {"underline": True})])
    _mk_para(doc, [
        ("工期：", None),
        ("_____天", {"underline": True}),
        ("（含日历天）", None),
    ])
    _mk_para(doc, [
        ("（九）", None),
        ("项目经理姓名：", None),
        ("__________", {"underline": True}),
    ])
    doc.save(path)
    return path


class TestUnderlineFidelity:
    def test_单run保真(self, tmp_path):
        tpl = _mk_template(str(tmp_path))
        doc = Document(tpl)
        p = doc.paragraphs[0]
        ms = list(fd.PH_RE.finditer(p.text or ""))
        assert fd._replace_in_para(p, ms[0].group(0), "淄博油库工程")
        doc.save(str(tmp_path / "out.docx"))
        doc2 = Document(str(tmp_path / "out.docx"))
        _assert_underline_in(doc2.paragraphs[0], "淄博油库工程")

    def test_跨run_中段下划线(self, tmp_path):
        tpl = _mk_template(str(tmp_path))
        doc = Document(tpl)
        p = doc.paragraphs[1]
        ms = list(fd.PH_RE.finditer(p.text or ""))
        assert fd._replace_in_para(p, ms[0].group(0), "240")
        doc.save(str(tmp_path / "out.docx"))
        doc2 = Document(str(tmp_path / "out.docx"))
        # 填入的「240」+ 紧跟 post「天」应在同一 rPr 下保留
        run_240 = next(r for r in doc2.paragraphs[1].runs if "240" in r.text)
        assert _has_underline(run_240), "跨 run 中段填入下划线丢失"
        # 「天」字必须仍存在（不被吞），且与「240」同一带下划线的 run
        assert "天" in run_240.text, f"post「天」字被吞：run={run_240.text!r}"

    def test_跨run_末段下划线(self, tmp_path):
        tpl = _mk_template(str(tmp_path))
        doc = Document(tpl)
        p = doc.paragraphs[2]
        ms = list(fd.PH_RE.finditer(p.text or ""))
        assert fd._replace_in_para(p, ms[0].group(0), "张三")
        doc.save(str(tmp_path / "out.docx"))
        doc2 = Document(str(tmp_path / "out.docx"))
        _assert_underline_in(doc2.paragraphs[2], "张三")

    def test_替换不改变段文本长度近似(self, tmp_path):
        # 跨 run 替换不应出现「天」字重复或「天」字丢失
        tpl = _mk_template(str(tmp_path))
        doc = Document(tpl)
        p = doc.paragraphs[1]
        ms = list(fd.PH_RE.finditer(p.text or ""))
        before = p.text
        fd._replace_in_para(p, ms[0].group(0), "240")
        after = p.text
        # before = "工期：_____天（含日历天）"
        # after  = "工期：240天（含日历天）"
        assert after.count("天") == before.count("天"), f"天字重复/丢失：{before!r} -> {after!r}"

    def test_未命中占位返回False(self, tmp_path):
        doc = Document()
        doc.add_paragraph("没有占位的段落")
        p = doc.paragraphs[0]
        assert fd._replace_in_para(p, "____", "x") is False
