"""对比 fill_docx 替换前后的下划线保真情况。

测试场景：
- 段1: "项目名称：____"（下划线占位 4 字符）
- 段2: "工期：_____天"（下划线占位 5 字符，紧跟"天"）
- 段3: "（九）工期：_______天"（带编号前缀 + 7 字符下划线）

对每个段统计：
- PH_RE 命中的下划线 run 的 w:u（underline）值（替换前）
- 替换后整段内 w:u 的存在情况（验证「填完下划线是否消失」）

修复目标：填入文字本身带下划线；非填入部分（如"天"字）保留原格式。
"""
import json, os, sys, tempfile
sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "scripts"))

from docx import Document
from docx.oxml.ns import qn

d = tempfile.mkdtemp()
tpl = os.path.join(d, "template.docx")
doc = Document()
doc.add_paragraph("项目名称：____")
doc.add_paragraph("工期：_____天")
doc.add_paragraph("（九）工期：_______天")
doc.save(tpl)

# 显式为「____」run 加下划线（模拟招标文件原样切片的实际格式）
doc2 = Document(tpl)
from lxml import etree
W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
A = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
for p in doc2.paragraphs:
    # 找下划线段：包含连续 2+ 个「_」的 run
    for r in p.runs:
        if "_" * 2 in r.text:
            rPr = r._element.find(qn("w:rPr"))
            if rPr is None:
                rPr = etree.SubElement(r._element, qn("w:rPr"))
                r._element.insert(0, rPr)
            # 加 w:u val="single"
            u = rPr.find(qn("w:u"))
            if u is None:
                u = etree.SubElement(rPr, qn("w:u"))
            u.set(qn("w:val"), "single")
doc2.save(tpl)

# 修复前的扫描：仅替换文本
doc_check = Document(tpl)
print("=== 修复前：仅 r.text 替换 ===")
for i, p in enumerate(doc_check.paragraphs):
    print(f"  段{i}: text={p.text!r}")
    for j, r in enumerate(p.runs):
        rPr = r._element.find(qn("w:rPr"))
        u = rPr.find(qn("w:u")) if rPr is not None else None
        print(f"    run{j}: text={r.text!r} has_underline={u is not None}")

# 调用原版 fill_docx
from fill_docx import PH_RE, replace_in_para
out_old = os.path.join(d, "old.docx")
fills = {"P0#0": "淄博油库工程", "P1#0": "240天", "P2#0": "46天"}
fills_data = {"fills": fills, "unconfirmed": 0}
fills_path = os.path.join(d, "fills.json")
json.dump(fills_data, open(fills_path, "w", encoding="utf-8"), ensure_ascii=False)

# 模拟旧 replace_in_para 跑一遍
doc_old = Document(tpl)
for pid, val in fills.items():
    pi = int(pid.split("#")[0][1:])
    k = int(pid.split("#")[1])
    p = doc_old.paragraphs[pi]
    ms = list(PH_RE.finditer(p.text or ""))
    if k < len(ms):
        replace_in_para(p, ms[k].group(0), str(val))
doc_old.save(out_old)

doc_old_chk = Document(out_old)
print("\n=== 旧版 fill_docx 替换后 ===")
for i, p in enumerate(doc_old_chk.paragraphs):
    print(f"  段{i}: text={p.text!r}")
    for j, r in enumerate(p.runs):
        rPr = r._element.find(qn("w:rPr"))
        u = rPr.find(qn("w:u")) if rPr is not None else None
        print(f"    run{j}: text={r.text!r} has_underline={u is not None}")
