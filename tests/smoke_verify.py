"""verify.py 冒烟：四轨道 + SKIP + D 短行判定。
直接 import verify.main 而不是 subprocess，避免 docx 依赖在不同 venv 不可用。
"""
import json, os, sys, tempfile
sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "scripts"))
import verify as vmod

d = tempfile.mkdtemp()
tender = os.path.join(d, "tender.json")
quotes = os.path.join(d, "quotes.json")
target = os.path.join(d, "draft.docx")
out_full = os.path.join(d, "report_full.md")
out_skip = os.path.join(d, "report_skip.md")

# tender：含一条目录短行 + 一条正文长行
json.dump({"disqualification_clauses": [
    {"clause": "实质性条款响应承诺"},
    {"clause": "商务报价excel版"}
]}, open(tender, "w", encoding="utf-8"), ensure_ascii=False)

# quotes：大小写与合计一致（应 PASS）
json.dump({"quotes": [
    {"label": "材料费", "amount": "10000", "amount_cn": "壹万元整"},
    {"label": "人工费", "amount": "5000", "amount_cn": "伍仟元整"},
    {"label": "总价", "amount": "15000", "amount_cn": "壹万伍仟元整"}
]}, open(quotes, "w", encoding="utf-8"), ensure_ascii=False)

from docx import Document
doc = Document()
doc.add_paragraph("目录")
doc.add_paragraph("实质性条款响应承诺")  # 短行
doc.add_heading("正文", 1)
doc.add_paragraph("商务报价excel版的内容：已按要求提交电子版商务报价Excel表")
doc.save(target)

# 构造 mock args 直接调 main()
class Args:
    pass

def run(args):
    import io
    from contextlib import redirect_stdout, redirect_stderr
    buf_out, buf_err = io.StringIO(), io.StringIO()
    try:
        with redirect_stdout(buf_out), redirect_stderr(buf_err):
            vmod.main() if False else _run_main_with_args(args)
        rc = 0
    except SystemExit as e:
        rc = e.code
    return rc, buf_out.getvalue(), buf_err.getvalue()

def _run_main_with_args(args):
    import argparse
    # monkey-patch argparse
    orig = argparse.ArgumentParser.parse_args
    def fake_parse(self_inner, *a, **kw):
        return args
    argparse.ArgumentParser.parse_args = fake_parse
    try:
        vmod.main()
    finally:
        argparse.ArgumentParser.parse_args = orig

# Case 1: 全参 → A PASS, B PASS（合计+大写）, C SKIP（无 template）, D PASS（两条都在长行/不存在短行）
# 注：实测"实质性条款响应承诺"在 docx 里就是短行，应判 short_only → WEAK
print("=== Case 1：全参 ===")
a = Args(); a.target = target; a.template = ""; a.quotes = quotes; a.tender = tender; a.out = out_full
rc, out, err = run(a)
print(f"rc={rc} | stdout={out.strip()}")
print(f"stderr: {err.strip()}")
report = open(out_full, encoding="utf-8").read()
for ln in report.splitlines():
    if "|" in ln and ("PASS" in ln or "FAIL" in ln or "WEAK" in ln or "SKIP" in ln):
        print("  ", ln)
print()

# Case 2: 仅 target → B/C/D 三项 SKIP
print("=== Case 2：仅 target ===")
a2 = Args(); a2.target = target; a2.template = ""; a2.quotes = ""; a2.tender = ""; a2.out = out_skip
rc2, out2, err2 = run(a2)
print(f"rc={rc2} | stdout={out2.strip()}")
print(f"stderr: {err2.strip()}")
report2 = open(out_skip, encoding="utf-8").read()
for ln in report2.splitlines():
    if "SKIP" in ln:
        print("  ", ln)
