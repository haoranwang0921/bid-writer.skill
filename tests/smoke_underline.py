"""验证 fill_docx v0.3.2 下划线保真 + 跨 run 修复。

构造 docx：占位段以「单 run 整段」 + 「跨 3 run」两种结构，验证下划线 run 的 rPr
在替换后保留。
"""
import json, os, sys, tempfile
sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "scripts"))

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree


def add_underline(rPr):
    """给 rPr 节点加 w:u val=single（幂等）。"""
    u = rPr.find(qn("w:u"))
    if u is None:
        u = etree.SubElement(rPr, qn("w:u"))
    u.set(qn("w:val"), "single")
    return rPr


def mk_para(doc, runs_spec):
    """runs_spec = [(text, {'underline': bool/fmt, 'bold': ...})]
    多个 tuple 即构造多 run 段落。"""
    p = doc.add_paragraph()
    for text, props in runs_spec:
        r = p.add_run(text)
        if props:
            rPr = r._element.find(qn("w:rPr"))
            if rPr is None:
                rPr = OxmlElement("w:rPr")
                r._element.insert(0, rPr)
            if props.get("underline"):
                add_underline(rPr)
            if props.get("bold"):
                b = OxmlElement("w:b")
                rPr.append(b)
    return p


d = tempfile.mkdtemp()
tpl = os.path.join(d, "template.docx")
doc = Document()

# 段 0：单 run 含下划线占位
mk_para(doc, [("项目名称：____", {"underline": True})])
# 段 1：跨 3 run（前/中/后），仅中间 run 带下划线
mk_para(doc, [
    ("工期：", None),
    ("_____天", {"underline": True}),  # 实际是「下划线段 + 天」，占位是「_____」
    ("（含日历天）", None),
])
# 段 2：跨 run 模拟 Word 自动拆字（前导空白 + 下划线占位 + 后缀）
mk_para(doc, [
    ("（九）", None),
    ("项目经理姓名：", None),
    ("__________", {"underline": True}),
])

doc.save(tpl)

# 用新 _replace_in_para 跑替换
from fill_docx import PH_RE, _replace_in_para
out = os.path.join(d, "filled.docx")
doc2 = Document(tpl)
for pi, (val, k) in enumerate([("淄博油库工程", 0), ("240", 0), ("张三", 0)]):
    p = doc2.paragraphs[pi]
    ms = list(PH_RE.finditer(p.text or ""))
    if k < len(ms):
        ok = _replace_in_para(p, ms[k].group(0), val)
        print(f"段{pi} 替换 {'成功' if ok else '失败'}: old={ms[k].group(0)!r} new={val!r}")
doc2.save(out)

# 校验：每个段填入文字的下划线 rPr 是否保留
doc3 = Document(out)
print("\n=== 替换后 run 状态 ===")
for i, p in enumerate(doc3.paragraphs):
    print(f"段{i}: text={p.text!r}")
    for j, r in enumerate(p.runs):
        rPr = r._element.find(qn("w:rPr"))
        u = rPr.find(qn("w:u")) if rPr is not None else None
        print(f"  run{j}: text={r.text!r} has_underline={u is not None}")

# 断言
def check_underline(para, must_underline_substr):
    """断言段落中包含 must_underline_substr 的 run 仍有下划线。"""
    for r in para.runs:
        if must_underline_substr in r.text:
            rPr = r._element.find(qn("w:rPr"))
            assert rPr is not None and rPr.find(qn("w:u")) is not None, \
                f"段 '{para.text}' 中「{r.text}」下划线丢失"
            return True
    assert False, f"段 '{para.text}' 中找不到「{must_underline_substr}」"

check_underline(doc3.paragraphs[0], "淄博油库工程")
check_underline(doc3.paragraphs[1], "240")
check_underline(doc3.paragraphs[2], "张三")
print("\n✓ 全部填入文字下划线保留")

# 跑原 fill_docx.py 主流程做端到端
print("\n=== 端到端：fill_docx.py 主流程 ===")
fills = {"P0#0": "淄博油库工程", "P1#0": "240", "P2#0": "张三"}
fills_path = os.path.join(d, "fills.json")
json.dump({"fills": fills}, open(fills_path, "w", encoding="utf-8"), ensure_ascii=False)

import subprocess
py = r"C:\Users\whr\.workbuddy\binaries\python\envs\default\Scripts\python.exe"
script = os.path.join(os.path.dirname(__file__), "..", "scripts", "fill_docx.py")
r = subprocess.run([py, script, tpl, fills_path, "-o", out], capture_output=True, text=True)
print(f"stdout: {r.stdout.strip()}")
print(f"stderr: {r.stderr.strip()}")
print(f"rc: {r.returncode}")

doc4 = Document(out)
for i, p in enumerate(doc4.paragraphs):
    for run in p.runs:
        if "淄博" in run.text or "240" in run.text or "张三" in run.text:
            rPr = run._element.find(qn("w:rPr"))
            assert rPr is not None and rPr.find(qn("w:u")) is not None, \
                f"端到端：段{i}「{run.text}」下划线丢失"
print("✓ 端到端下划线保留")
